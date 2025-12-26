import streamlit as st
from gusregon import GUS
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime, timedelta, timezone
import io

# --- 1. KONFIGURACJA STRONY (Musi być pierwsza) ---
st.set_page_config(page_title="Elite Waste System", page_icon="Nowe Logo.png")

# --- 2. UKRYWANIE ELEMENTÓW STREAMLIT (Stylizacja) ---
hide_streamlit_style = """
            <style>
            #MainMenu {visibility: hidden;}
            footer {visibility: hidden;}
            header {visibility: hidden;}
            [data-testid="stToolbar"] {visibility: hidden !important;}
            </style>
            """
st.markdown(hide_streamlit_style, unsafe_allow_html=True)

# --- 3. LOGO I NAGŁÓWEK ---
try:
    st.image("Nowe Logo.png", width=300)
except Exception:
    st.warning("⚠️ Nie znaleziono pliku 'Nowe Logo.png'. Upewnij się, że plik jest wgranym na GitHub.")

st.title("♻️ Generator Pełnomocnictw BDO")
st.markdown("### Elite Waste")
st.info("Wpisz NIP klienta poniżej. System obsługuje zarówno spółki (KRS) jak i JDG (CEIDG).")

# --- 4. POBIERANIE KLUCZA ---
try:
    api_key = st.secrets["GUS_KEY"]
except Exception as e:
    st.error("⚠️ Błąd konfiguracji! Nie znaleziono klucza GUS_KEY w zakładce Secrets.")
    st.stop()

# --- 5. INTELIGENTNE WYCIĄGANIE DANYCH ---
def wyciagnij_dane_smart(dane):
    """
    Funkcja mapuje nazwy pól z GUS na podstawie Twoich zrzutów ekranu.
    """
    # 1. NAZWA
    nazwa = dane.get('nazwa', '')
    
    # 2. REGON (GUS dla JDG wrzuca go do 'regon9')
    regon = dane.get('regon') or dane.get('regon9') or ""

    # 3. MIEJSCOWOŚĆ
    miasto = (dane.get('adsiedzmiejscowosc_nazwa') 
              or dane.get('siedzibamiejscowosc_nazwa') 
              or dane.get('miejscowosc') 
              or "")
    
    # 4. ULICA
    ulica = (dane.get('adsiedzulica_nazwa') 
             or dane.get('siedzibaulica_nazwa') 
             or dane.get('ulica') 
             or "")
    
    # 5. NUMERY
    nr_domu = (dane.get('adsiedznumernieruchomosci') 
               or dane.get('nr_nieruchomosci') 
               or "")
    
    nr_lokalu = (dane.get('adsiedznumerlokalu') 
                 or dane.get('nr_lokalu') 
                 or "")
    
    # 6. KOD POCZTOWY
    kod = (dane.get('adsiedzkodpocztowy') 
           or dane.get('kod_pocztowy') 
           or "")
    # Formatowanie kodu (dodanie myślnika jeśli go nie ma)
    if kod and len(kod) == 5 and '-' not in kod:
        kod = f"{kod[:2]}-{kod[2:]}"

    # 7. WOJEWÓDZTWO
    woj = (dane.get('adsiedzwojewodztwo_nazwa') 
           or dane.get('wojewodztwo') 
           or "")

    return {
        'nazwa': nazwa,
        'miasto': miasto,
        'ulica': ulica,
        'nr_domu': nr_domu,
        'nr_lokalu': nr_lokalu,
        'kod': kod,
        'wojewodztwo': woj,
        'regon': regon
    }

# --- 6. GENERATOR DOKUMENTU ---
def generuj_word(info, nip_raw):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)

    # NAPRAWA CZASU: Wymuszenie strefy czasowej UTC+1 (Polska Zima)
    strefa_pl = timezone(timedelta(hours=1))
    data_dzis = datetime.now(strefa_pl).strftime("%d.%m.%Y")

    # Data
    p = doc.add_paragraph(f"Łódź, dnia {data_dzis} r.")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Mocodawca
    doc.add_paragraph("\nMocodawca").runs[0].bold = True
    
    # Budowanie adresu
    adres_string = ""
    ulica_czysta = info['ulica']
    
    if ulica_czysta:
        if "ul." in ulica_czysta.lower():
             adres_string += f"{ulica_czysta} {info['nr_domu']}"
        else:
             adres_string += f"ul. {ulica_czysta} {info['nr_domu']}"
    else:
        adres_string += f"{info['nr_domu']}"
        
    if info['nr_lokalu']:
        adres_string += f"/{info['nr_lokalu']}"
    
    adres_string += f", {info['kod']} {info['miasto']}"

    # Nagłówek danych firmy
    doc.add_paragraph(info['nazwa'].upper()) 
    doc.add_paragraph(adres_string)
    doc.add_paragraph(f"NIP: {nip_raw}")
    doc.add_paragraph(f"REGON: {info['regon']}")

    # Tytuł
    tytul = doc.add_paragraph("\nPEŁNOMOCNICTWO")
    tytul.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tytul.runs[0].bold = True
    tytul.runs[0].font.size = Pt(14)

    # Województwo
    woj_text = info['wojewodztwo'].lower()
    mapa_woj = {
        'łódzkie': 'Łódzkiego', 'mazowieckie': 'Mazowieckiego', 'wielkopolskie': 'Wielkopolskiego',
        'małopolskie': 'Małopolskiego', 'śląskie': 'Śląskiego', 'pomorskie': 'Pomorskiego',
        'dolnośląskie': 'Dolnośląskiego', 'podkarpackie': 'Podkarpackiego', 'lubelskie': 'Lubelskiego',
        'kujawsko-pomorskie': 'Kujawsko-Pomorskiego', 'zachodniopomorskie': 'Zachodniopomorskiego',
        'warmińsko-mazurskie': 'Warmińsko-Mazurskiego', 'świętokrzyskie': 'Świętokrzyskiego',
        'podlaskie': 'Podlaskiego', 'opolskie': 'Opolskiego', 'lubuskie': 'Lubuskiego'
    }
    
    urzad_wojewodztwo = mapa_woj.get(woj_text, woj_text.capitalize())
    if not urzad_wojewodztwo:
         urzad_wojewodztwo = "........................................"

    # Treść
    czesc_1 = f"Działając w imieniu {info['nazwa']} z siedzibą w {info['miasto']}, {adres_string}, posiadając prawo reprezentacji tego podmiotu w zakresie ustanawiania pełnomocnictw, upoważniam Pana Pawła Bolimowskiego oraz Pana Patryka Kosteckiego do samodzielnej reprezentacji "
    czesc_2 = f"{info['nazwa']} przed Urzędem Marszałkowskim Województwa {urzad_wojewodztwo} w następujących sprawach załatwianych za pośrednictwem indywidualnego konta w Bazie danych o produktach i opakowaniach oraz o gospodarce odpadami (BDO):\n"
    
    tekst = czesc_1 + czesc_2
    p_main = doc.add_paragraph(tekst)
    p_main.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Lista czynności
    punkty = [
        "złożenia wniosku o wpis do rejestru na wniosek zgodnie z art. 50 ustawy o odpadach;",
        "wyznaczania upoważnionych użytkowników zgodnie z art. 79 ust. 7 ustawy o odpadach;",
        "złożenia wniosku o zmianę wpisu w rejestrze zgodnie z art. 59 ustawy o odpadach;",
        "złożenia wniosku o wykreślenie z rejestru zgodnie z art. 60 ustawy o odpadach;",
        "prowadzenia ewidencji odpadów zgodnie z art. 66 i nast. ustawy o odpadach;",
        "prowadzenia sprawozdawczości zgodnie z art. 73 i nast. ustawy o odpadach."
    ]
    for punkt in punkty:
        p = doc.add_paragraph(f"- {punkt}")
        p.paragraph_format.left_indent = Cm(1)

    # Stopka
    doc.add_paragraph(f"\nPełnomocnictwo ustanawia się od dnia {data_dzis} r. do odwołania.")
    doc.add_paragraph(
        "Odwołanie pełnomocnictwa nie powoduje unieważnienia czynności wykonanych przez upoważnioną osobę "
        "ani konsekwencji tych czynności, jeżeli czynność miała miejsce przed poinformowaniem organu właściwego o cofnięciu pełnomocnictwa."
    )
    doc.add_paragraph("\n\n..................................................................")
    doc.add_paragraph("(Czytelny podpis oraz pieczątka Mocodawcy)")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 7. LOGIKA APLIKACJI ---
nip_input = st.text_input("Podaj NIP (sam numer, bez kresek):", max_chars=10)

if st.button("🔍 Znajdź firmę i generuj dokument"):
    if not nip_input:
        st.warning("Proszę wpisać NIP.")
    else:
        try:
            gus = GUS(api_key=api_key)
            dane_raw = gus.search(nip=nip_input)
            
            # Inteligentne wyciąganie danych
            info = wyciagnij_dane_smart(dane_raw)
            
            if not info['miasto']:
                st.error("Nie udało się pobrać adresu. Poniżej surowe dane:")
                with st.expander("Zobacz surowe dane"):
                    st.json(dane_raw)
            else:
                st.success(f"Znaleziono: **{info['nazwa']}**")
                
                # Generowanie pliku
                plik_word = generuj_word(info, nip_input)
                
                st.markdown("### 👇 Pobierz gotowy plik:")
                st.download_button(
                    label="📥 POBIERZ PEŁNOMOCNICTWO (DOCX)",
                    data=plik_word,
                    file_name=f"Pelnomocnictwo_BDO_{nip_input}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
        except Exception as e:
            st.error(f"Wystąpił błąd. (Szczegóły: {e})")
